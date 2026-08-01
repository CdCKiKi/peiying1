"""
生成测试文件和测试指南
用法: cd test_data && ..\\apps\\api\\venv\\Scripts\\python.exe generate_tests.py
"""
import os
import sys

# 确保依赖可用
try:
    import fitz  # PyMuPDF
    from PIL import Image, ImageDraw, ImageFont
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as e:
    print(f"缺少依赖: {e}")
    print("请运行: pip install PyMuPDF Pillow python-docx")
    sys.exit(1)

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def make_pdf(filename: str, title: str, lines: list[str]) -> str:
    """生成 PDF 文件"""
    path = os.path.join(OUTPUT_DIR, filename)
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)  # A4

    y = 60
    # 标题
    page.insert_text(fitz.Point(50, y), title, fontsize=18, fontname="china-s")
    y += 30
    page.insert_text(fitz.Point(50, y), "─" * 70, fontsize=10, fontname="china-s")
    y += 30

    for line in lines:
        page.insert_text(fitz.Point(50, y), line, fontsize=12, fontname="china-s")
        y += 22

    doc.save(path)
    doc.close()
    print(f"  已生成: {filename}")
    return path


def make_image(filename: str, lines: list[str]) -> str:
    """生成模拟扫描件图片"""
    path = os.path.join(OUTPUT_DIR, filename)
    img = Image.new("RGB", (800, 500), "white")
    draw = ImageDraw.Draw(img)

    # 模拟纸张效果
    for x in range(800):
        for y in range(500):
            r, g, b = img.getpixel((x, y))
            noise = int((x * 7 + y * 13) % 20)
            if noise < 15:
                img.putpixel((x, y), (255 - noise % 5, 255 - noise % 5, 250 - noise % 8))

    draw = ImageDraw.Draw(img)
    y = 30
    for line in lines:
        draw.text((40, y), line, fill="black")
        y += 28

    # 加一个框模拟扫描边界
    draw.rectangle([15, 15, 785, 485], outline="gray", width=2)

    img.save(path)
    print(f"  已生成: {filename}")
    return path


# ===== 1. 租务通知 PDF =====
make_pdf("test_01_rental_notice.pdf", "俊傑花園租金通知", [
    "俊傑花園租金通知",
    "",
    "單位：A座 8樓 B室",
    "租戶：陳先生",
    "月份：2026年8月",
    "租金：HK$ 18,500",
    "繳付限期：2026年8月31日",
    "",
    "請於限期前將租金存入以下銀行戶口：",
    "匯豐銀行 123-456789-001",
    "培英中學",
    "",
    "如有查詢，請聯絡校務處 Tommy。",
    "電話：2345 6789",
    "",
    "培英中學校務處",
    "2026年7月25日",
])

# ===== 2. 財務收據 PDF =====
make_pdf("test_02_receipt.pdf", "維修費收據", [
    "冷氣維修服務 收據",
    "",
    "收據編號：INV-2026-0722",
    "日期：2026年7月22日",
    "",
    "項目：俊傑花園 A座 冷氣系統維修",
    "供應商：香港冷氣工程有限公司",
    "金額：HK$ 2,300",
    "",
    "明細：",
    "  冷氣壓縮機更換：HK$ 1,800",
    "  雪種補充：HK$ 500",
    "",
    "付款方式：銀行轉帳",
    "經辦人：Tommy",
    "",
    "培英中學財務部",
])

# ===== 3. 教育局通告 PDF =====
make_pdf("test_03_edb_notice.pdf", "教育局通告", [
    "教育局通告",
    "",
    "通告編號：EDB/CD/2026/045",
    "日期：2026年7月18日",
    "",
    "主旨：2026/27學年學校暑期校園安全指引",
    "",
    "敬啟者：",
    "",
    "為確保暑假期間校園安全，請各學校注意以下事項：",
    "",
    "1. 校園保安：暑假期間應確保校園出入口妥善鎖閉，",
    "   安排定時巡邏。",
    "",
    "2. 防火安全：檢查滅火設備是否在有效期內。",
    "",
    "3. 水電檢查：離開校園前關閉非必要的電源和水源。",
    "",
    "如有任何疑問，請聯絡分區教育主任。",
    "",
    "教育局學校發展分部",
    "2026年7月18日",
])

# ===== 4. 會議記錄 PDF =====
make_pdf("test_04_meeting.pdf", "校務會議記錄", [
    "培英中學 校務會議記錄",
    "",
    "會議日期：2026年7月15日",
    "會議時間：下午 2:30 - 4:00",
    "會議地點：校務會議室",
    "主持：李校長",
    "出席：陳副校長、王主任、張老師、Tommy",
    "",
    "議程：",
    "1. 2026/27學年開學準備工作",
    "2. 暑期校園維修工程進度",
    "3. 俊傑花園租務事宜",
    "4. 其他事項",
    "",
    "決議：",
    "1. 開學日定於2026年9月1日",
    "2. 暑期維修工程預算追加至 HK$ 80,000",
    "3. 租金調整方案將於下次會議討論",
    "",
    "下次會議：2026年8月12日",
])

# ===== 5. 人事合同 DOCX =====
docx_path = os.path.join(OUTPUT_DIR, "test_05_hr_contract.docx")
doc = Document()
doc.styles["Normal"].font.size = Pt(11)
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("培英中學 員工合約")
run.bold = True
run.font.size = Pt(16)
doc.add_paragraph("")
for line in [
    "合約編號：HR-2026-0032",
    "日期：2026年7月10日",
    "",
    "員工姓名：陳小華",
    "職位：教學助理",
    "薪金：HK$ 16,500 / 月",
    "入職日期：2026年9月1日",
    "合約期：2026年9月1日至2027年8月31日",
    "",
    "假期：按學校校曆表安排",
    "福利：強積金、醫療保險",
    "",
    "培英中學人事部",
]:
    doc.add_paragraph(line)
doc.save(docx_path)
print(f"  已生成: test_05_hr_contract.docx")

# ===== 6. 扫描图片 PNG =====
make_image("test_06_scan_image.png", [
    "俊傑花園 車位租金通知",
    "",
    "車位：車位03",
    "租戶：李太太",
    "月份：2026年8月",
    "租金：HK$ 2,500",
    "繳付限期：2026年8月31日",
    "",
    "請準時繳交。",
    "培英中學校務處",
])

# ===== 生成测试指南 Word 文档 =====
print("\n生成测试指南...")
guide_path = os.path.join(OUTPUT_DIR, "培英AI平台_测试指南.docx")
guide = Document()

# 样式
style = guide.styles["Normal"]
style.font.size = Pt(11)
style.font.name = "Microsoft YaHei"

def add_title(text: str):
    t = guide.add_heading(text, level=0)

def add_h1(text: str):
    guide.add_heading(text, level=1)

def add_h2(text: str):
    guide.add_heading(text, level=2)

def add_body(text: str):
    guide.add_paragraph(text)

def add_step(num: int, text: str):
    guide.add_paragraph(f"{num}. {text}")

def add_table(headers: list[str], rows: list[list[str]]):
    table = guide.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Light List Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r + 1].cells[c].text = val
    guide.add_paragraph("")

add_title("培英 AI 行政平台 - 测试指南")

guide.add_paragraph("版本：1.0 | 日期：2026年7月27日 | 测试环境：Windows 11, Python 3.12, Node.js 20")
guide.add_paragraph("")

# ===== 第一章：环境准备 =====
add_h1("一、测试环境准备")

add_h2("1.1 启动系统")
add_step(1, "确保 .env 文件中已配置 LLM API Key（DeepSeek 或其他）")
add_step(2, "双击项目根目录下的 start.bat")
add_step(3, "等待两个窗口分别显示：")
add_body("   - Backend API：Uvicorn running on http://0.0.0.0:8000")
add_body("   - Frontend Web：Ready in X.Xs (http://localhost:3000)")
add_step(4, "如果后端启动失败，检查 .env 中的 LLM_API_KEY 和数据库路径")

add_h2("1.2 测试账号")
add_table(["用户名", "密码", "角色", "说明"], [
    ["tommy", "tommy123", "校務處書記", "主要测试账号，有文件归档和租务管理权限"],
    ["admin", "admin123", "系統管理員", "管理员账号"],
])

add_h2("1.3 测试文件")
add_body("test_data 目录下已生成 6 份测试文件：")
add_table(["文件名", "类型",  "预期分类", "测试目的"], [
    ["test_01_rental_notice.pdf", "PDF", "租務", "租务通知 OCR + AI 分类"],
    ["test_02_receipt.pdf", "PDF", "財務", "财务收据识别"],
    ["test_03_edb_notice.pdf", "PDF", "教育局通告", "教育局通告识别"],
    ["test_04_meeting.pdf", "PDF", "會議", "会议记录识别"],
    ["test_05_hr_contract.docx", "DOCX", "人事", "人事合同分类"],
    ["test_06_scan_image.png", "PNG", "租務", "图片文件 OCR + 预览"],
])

# ===== 第二章：功能测试 =====
add_h1("二、功能测试清单")

add_h2("2.1 登入功能")
add_step(1, "打开 http://localhost:3000")
add_step(2, "用 tommy / tommy123 登录")
add_step(3, "验证：跳转到 Tommy 工作台总览页")
add_step(4, "验证：左侧边栏显示「總覽」「文件歸檔」「租務提醒」")
add_step(5, "测试错误密码：确认显示错误提示")
add_step(6, "点击右上角设置图标，测试修改密码功能")

add_h2("2.2 文件上传与智能归档（核心流程）")
add_step(1, "点击侧边栏「文件歸檔」进入归档页")
add_step(2, "点击「上傳文件」按钮")
add_step(3, "拖拽或选择 test_01_rental_notice.pdf")
add_step(4, "验证：上传后文档状态显示「待處理」→ 几秒后自动变为「待確認」")
add_step(5, "验证：OCR 原文显示在右侧预览面板")
add_step(6, "验证：AI 分类结果（分类应为「租務」，金额 HK$ 18,500，到期日 2026-08-31）")
add_step(7, "编辑分类/文件名/金额/日期（模拟人工审核修改）")
add_step(8, "点击「確認保存」→ 状态变为「已確認」")
add_step(9, "点击「確認並歸檔」→ 状态变为「已歸檔」")
add_body("")
add_body("重复上传 test_02 ~ test_06，验证各自分类准确率。")

add_h2("2.3 文件预览")
add_step(1, "上传 test_06_scan_image.png（或选择已有的图片文件）")
add_step(2, "在文件列表中点击「預覽」")
add_step(3, "验证：图片预览弹窗中直接显示图片内容")
add_step(4, "上传 test_01_rental_notice.pdf 后点击预览")
add_step(5, "验证：PDF 在 iframe 中正常预览")
add_step(6, "点击「在新窗口打開」→ 新标签页中打开文件")
add_step(7, "点击「下載文件」→ 文件下载到本地")

add_h2("2.4 搜索功能")
add_step(1, "在归档页点击「更多操作」")
add_step(2, "在搜索框中输入「租金」→ 验证返回包含租金关键词的文件")
add_step(3, "输入 OCR 原文中的内容「俊傑花園」→ 验证全文搜索")
add_step(4, "按分类筛选：选择「租務」→ 验证只显示租务类文件")
add_step(5, "按状态筛选：选择「已歸檔」→ 验证只显示已归档文件")
add_step(6, "使用日期范围筛选：设置起止日期 → 验证显示该时段内的文件")
add_step(7, "点击「清除篩選」→ 验证恢复显示全部")

add_h2("2.5 分页导航")
add_step(1, "上传足够多的文件（至少 21 份以触发第二页）")
add_step(2, "验证：页面底部出现分页组件")
add_step(3, "点击「下一頁」→ 验证跳转到第 2 页")
add_step(4, "点击页码数字 → 验证直接跳转")
add_step(5, "在跳转输入框输入页码后按 Enter → 验证跳转")
add_step(6, "切换每页条数（10/20/50/100）→ 验证生效")

add_h2("2.6 批量操作")
add_step(1, "在归档页勾选多份文件左侧的复选框")
add_step(2, "验证：顶部出现批量操作栏「批量确认」「批量归档」「批量重试」")
add_step(3, "点击「批量確認」→ 验证多份文件同时变为已确认")
add_step(4, "勾选已确认文件 → 点击「批量歸檔」→ 验证批量归档成功")
add_step(5, "点击「全选」复选框 → 验证全选/取消全选")
add_step(6, "点击「取消選擇」→ 验证清除所有勾选")

add_h2("2.7 异常处理与重试")
add_step(1, "模拟异常场景：暂时断网或修改 OCR Backend 为 mock，查看异常处理")
add_step(2, "对异常状态文件点击「重試」→ 验证重新执行 OCR+AI")
add_step(3, "使用批量重试：勾选多个异常文件 → 点击「批量重試」")
add_step(4, "验证：重试成功后状态变为「待確認」")

add_h2("2.8 撤销操作")
add_step(1, "对已确认文件点击「撤銷」→ 验证状态回退为「待確認」")
add_step(2, "对已归档文件点击「撤銷」→ 验证状态回退为「已確認」")
add_step(3, "确认撤销只能回退一步")

add_h2("2.9 租务管理 - 单位列表")
add_step(1, "点击侧边栏「租務提醒」进入租务管理页")
add_step(2, "验证：显示 8 个住宅单位 + 9 个车位")
add_step(3, "验证：30 天内到期单位显示红色「剩 X 天」")
add_step(4, "验证：逾期缴费单位行显示红色背景 + 「逾期 X 天」")
add_step(5, "验证：分页组件正常显示")

add_h2("2.10 租务管理 - 租约操作")
add_step(1, "点击某单位的「查看」按钮 → 弹出详情弹窗")
add_step(2, "验证：弹窗显示物业名称、租户、月租、租约期间、电邮")
add_step(3, "验证：弹窗中的「租約記錄」表格显示历史租约")
add_step(4, "验证：弹窗中的「繳費記錄」表格显示缴费状态")
add_step(5, "点击租约的「編輯」→ 修改租约信息 → 确认保存")
add_step(6, "点击「新增租約」→ 选择「新建單位」→ 填写信息 → 确认")
add_step(7, "测试续租：选择单位为已有单位 → 填写新租约信息 → 确认")
add_step(8, "验证：旧租约自动标记为「已到期」")

add_h2("2.11 租务管理 - 缴费管理")
add_step(1, "点击「記錄繳費」→ 填写金额和截止日 → 确认")
add_step(2, "验证：缴费记录出现在单位详情的「繳費記錄」中")
add_step(3, "对待缴记录点击「收款」→ 验证状态变为已缴")
add_step(4, "对已缴记录点击「撤銷」→ 验证状态恢复为待缴")
add_step(5, "点击「發送繳費提醒」（需配置 SMTP）→ 验证审计日志记录")

add_h2("2.12 租约文件生成")
add_step(1, "在单位详情弹窗中点击「生成租約文件」")
add_step(2, "验证：浏览器自动下载 .docx 格式租约文件")
add_step(3, "打开下载的租约文件 → 验证内容正确（单位、租户、租期、金额）")

add_h2("2.13 仪表板总览")
add_step(1, "点击侧边栏「總覽」回到总览页")
add_step(2, "验证：4 个统计卡片显示正确数据")
add_step(3, "验证：饼图显示分类分布（上传多份不同类型的文件后查看）")
add_step(4, "验证：柱状图显示近 6 个月上传趋势")
add_step(5, "验证：租约到期预警卡片显示即将到期单位")
add_step(6, "验证：缴费逾期追踪卡片显示逾期缴费")
add_step(7, "验证：底部系统状态面板显示数据库/磁盘/错误/版本信息")
add_step(8, "点击任意统计卡 → 验证跳转到对应功能页")

add_h2("2.14 快速入口")
add_step(1, "在总览页点击「進入歸檔」→ 验证跳转到归档页")
add_step(2, "在总览页点击「進入租務」→ 验证跳转到租务页")

add_h2("2.15 审计日志")
add_step(1, "在归档页点击某文件的「查看詳情」")
add_step(2, "验证：弹窗底部「審計記錄」显示完整的操作历史")
add_step(3, "验证：记录包含时间、操作人、操作内容")

# ===== 第三章：验证标准 =====
add_h1("三、各功能验证标准")

add_table(["功能模块", "验证标准", "通过条件"], [
    ["登录", "输入正确账号密码后跳转到总览", "跳转成功"],
    ["文件上传", "上传后列表出现新文件", "文件出现在列表中"],
    ["异步OCR", "上传后3-6秒自动完成OCR+分类", "状态从待處理变为待確認"],
    ["AI分类准确率", "租务文件分类为「租務」", "分类正确（可人工修正）"],
    ["图片预览", "JPG/PNG 文件可在弹窗中直接查看", "图片显示正常"],
    ["PDF预览", "PDF 文件在 iframe 中显示", "PDF 显示正常"],
    ["全文搜索", "搜索 OCR 原文中的词能匹配到文件", "返回正确结果"],
    ["日期筛选", "设置日期范围后只显示该时段文件", "筛选正确"],
    ["分页", "文件超过每页条数时出现分页", "翻页功能正常"],
    ["批量确认", "勾选多份文件后批量确认", "所有选中文件状态变更"],
    ["批量归档", "勾选多份文件后批量归档", "所有选中文件状态变更"],
    ["异常重试", "异常文件重试后恢复", "状态恢复为待確認"],
    ["撤销操作", "确认/归档后可撤销回退", "状态正确回退"],
    ["租约新增", "新增租约后单位信息更新", "租户/租期更新正确"],
    ["缴费记录", "记录缴费后状态更新", "状态显示已缴"],
    ["生成租约", "点击生成后下载 .docx 文件", "文件下载成功"],
    ["仪表板图表", "饼图/柱状图显示正确数据", "图表渲染正常"],
    ["系统状态", "状态面板显示数据库/磁盘信息", "信息显示正确"],
])

# ===== 第四章：已知限制 =====
add_h1("四、已知限制与注意事项")
add_body("1. 邮件功能需要配置有效的 SMTP 服务器才能真实发送，否则仅记录审计日志。")
add_body("2. OCR 准确率取决于讯飞服务的可用性和文件清晰度，拍照件精度低于扫描件。")
add_body("3. AI 分类在 LLM API 不可用时会自动回退到关键词匹配模式。")
add_body("4. SQLite 数据库不支持多用户并发写入，生产环境建议切换 PostgreSQL。")
add_body("5. 首次启动会自动创建数据库和种子数据（仅当数据库为空时）。")
add_body("6. 讯飞 OCR 仅支持单页识别；多页 PDF 会分页处理。")

# ===== 第五章：故障排查 =====
add_h1("五、常见问题排查")

add_table(["问题", "可能原因", "解决方案"], [
    ["后端启动报错", "venv 路径不匹配", "使用 python -m uvicorn 代替 uvicorn"],
    ["前端 Failed to fetch", "后端未启动或端口冲突", "确认后端 8000 端口正常"],
    ["数据库表不存在", "Alembic 迁移未执行", "删除 school_ai_dev.db 后重启"],
    ["OCR 返回空文本", "讯飞 API 密钥过期或超限", "检查 XFYUN 配置或切换 mock"],
    ["AI 分类不准确", "LLM API Key 无效", "检查 LLM_API_KEY 是否正确"],
    ["搜索无结果", "FTS5 索引未初始化", "重启应用自动重建索引"],
    ["中文乱码", "控制台编码问题", "在 cmd 中执行 chcp 65001"],
])

guide.save(guide_path)
print(f"\n测试指南已生成: {guide_path}")
print(f"\n所有文件位于: {OUTPUT_DIR}")
print("测试文件列表:")
for f in sorted(os.listdir(OUTPUT_DIR)):
    if f != "generate_tests.py":
        size = os.path.getsize(os.path.join(OUTPUT_DIR, f))
        print(f"  {f} ({size:,} bytes)")
