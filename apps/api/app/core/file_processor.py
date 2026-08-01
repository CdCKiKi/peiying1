"""
文件预处理模块 - 支持 PDF、DOCX、图片格式的处理
"""
import os
import tempfile
import logging
from typing import List

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str) -> str:
    """从docx文件中提取文本"""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join([cell.text for cell in row.cells])
                if row_text.strip():
                    tables_text.append(row_text)
        return "\n".join(paragraphs + tables_text)
    except ImportError:
        logger.warning("python-docx 未安装，无法提取docx文本")
        return ""
    except Exception as e:
        logger.error(f"提取docx文本失败: {str(e)}")
        return ""


def extract_text_from_pdf(file_path: str) -> str:
    """从PDF文件中提取文本（使用PyMuPDF）"""
    try:
        import fitz
        doc = fitz.open(file_path)
        texts = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                texts.append(text)
        doc.close()
        return "\n\n".join(texts)
    except ImportError:
        logger.warning("PyMuPDF (fitz) 未安装，无法提取PDF文本")
        return ""
    except Exception as e:
        logger.error(f"提取PDF文本失败: {str(e)}")
        return ""


def pdf_to_images(file_path: str) -> List[str]:
    """将PDF转换为图片，返回图片路径列表"""
    try:
        from pdf2image import convert_from_path
        
        output_dir = tempfile.mkdtemp()
        images = convert_from_path(file_path)
        
        image_paths = []
        for i, image in enumerate(images):
            image_path = os.path.join(output_dir, f"page_{i+1}.png")
            image.save(image_path, "PNG")
            image_paths.append(image_path)
        
        logger.info(f"PDF转换完成，共 {len(image_paths)} 页")
        return image_paths
    except ImportError:
        logger.warning("pdf2image 未安装，无法转换PDF")
        return []
    except Exception as e:
        logger.error(f"PDF转换失败: {str(e)}")
        return []


def docx_to_images(file_path: str) -> List[str]:
    """将docx转换为图片，返回图片路径列表"""
    try:
        from docx import Document
        from PIL import Image as PILImage
        from docx.shared import Inches
        
        output_dir = tempfile.mkdtemp()
        doc = Document(file_path)
        
        image_paths = []
        for i, element in enumerate(doc.element.body):
            temp_doc = Document()
            temp_doc.element.body.append(element)
            
            temp_path = os.path.join(output_dir, f"temp_{i}.docx")
            temp_doc.save(temp_path)
            
            try:
                from pdf2image import convert_from_path
                temp_pdf = temp_path.replace(".docx", ".pdf")
                
                import subprocess
                result = subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "pdf", temp_path, "--outdir", output_dir],
                    capture_output=True,
                    timeout=60
                )
                
                if os.path.exists(temp_pdf):
                    images = convert_from_path(temp_pdf)
                    for j, img in enumerate(images):
                        img_path = os.path.join(output_dir, f"page_{i}_{j}.png")
                        img.save(img_path, "PNG")
                        image_paths.append(img_path)
            except Exception as e:
                logger.warning(f"转换docx页失败: {str(e)}")
        
        if not image_paths:
            text = extract_text_from_docx(file_path)
            if text:
                text_path = os.path.join(output_dir, "extracted_text.txt")
                with open(text_path, "w", encoding="utf-8") as f:
                    f.write(text)
                image_paths = [text_path]
        
        return image_paths
    except Exception as e:
        logger.error(f"docx转换失败: {str(e)}")
        return []


def get_file_type(file_path: str) -> str:
    """获取文件类型"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext in [".pdf"]:
        return "pdf"
    elif ext in [".docx", ".doc"]:
        return "docx"
    elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif"]:
        return "image"
    else:
        return "unknown"


def process_file_for_ocr(file_path: str) -> List[str]:
    """
    处理文件为OCR可用格式
    :param file_path: 文件路径
    :return: 图片路径列表（如果是文本文件则返回文本路径列表）
    """
    file_type = get_file_type(file_path)
    
    if file_type == "image":
        return [file_path]
    
    elif file_type == "pdf":
        text = extract_text_from_pdf(file_path)
        if text and len(text.strip()) > 10:
            output_dir = tempfile.mkdtemp()
            text_path = os.path.join(output_dir, "extracted_text.txt")
            with open(text_path, "w", encoding="utf-8") as f:
                f.write(text)
            logger.info(f"PDF文本提取成功，长度: {len(text)}")
            return [text_path]
        else:
            image_paths = pdf_to_images(file_path)
            if image_paths:
                return image_paths
            else:
                logger.warning("PDF转换失败且文本提取也失败")
                return []
    
    elif file_type == "docx":
        text = extract_text_from_docx(file_path)
        if text and len(text.strip()) > 10:
            output_dir = tempfile.mkdtemp()
            text_path = os.path.join(output_dir, "extracted_text.txt")
            with open(text_path, "w", encoding="utf-8") as f:
                f.write(text)
            return [text_path]
        else:
            return docx_to_images(file_path)
    
    else:
        logger.warning(f"不支持的文件类型: {file_type}")
        return []
