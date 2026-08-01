"""
租约文件生成器
使用 python-docx 生成正式租约文件
"""
import os
import logging
from datetime import date, datetime
from decimal import Decimal
from typing import Optional, Tuple
from io import BytesIO

from app.core.config import settings

logger = logging.getLogger(__name__)

# 租约模板常量
LEASE_TEMPLATE = """培英中學 - 俊傑花園租賃合約

本合約由以下雙方於 {lease_date} 共同簽訂：

出租方（甲方）：培英中學
地址：香港
電話：
傳真：

承租方（乙方）：{tenant_name}
身份證號碼：
聯絡電話：
通訊地址：

甲乙雙方經協商一致，就下列物業之租賃達成如下協議：

第一條 租賃物業
甲方同意將位於 俊傑花園 {unit_number} （以下簡稱「該物業」）出租予乙方作住宅用途使用。

第二條 租賃期限
1. 租賃期由 {lease_start} 至 {lease_end}，共 {lease_duration}。
2. 租賃期滿後，如乙方需續租，應於租賃期滿前兩個月以書面形式向甲方提出申請。

第三條 租金及付款方式
1. 每月租金為 港幣 HK$ {monthly_rent} 元正。
2. 乙方須於每月第 5 日前繳付當月租金。
3. 逾期繳付租金者，甲方有權按日加收應繳金額 0.1% 之滯納金。

第四條 押金
1. 乙方須於簽訂本合約時向甲方繳付相當於兩個月租金之押金，即 HK$ {deposit} 元正。
2. 租賃期滿且乙方已履行本合約全部義務後，甲方應於 30 日內將押金無息退還乙方。

第五條 物業使用
1. 乙方須妥善使用及維護該物業，不得對該物業進行任何結構性改動。
2. 乙方不得將該物業轉租或分租予任何第三方。
3. 未經甲方書面同意，乙方不得在該物業內進行任何非法活動。

第六條 費用承擔
1. 租賃期間，該物業之水費、電費、煤氣費等由乙方自行承擔。
2. 該物業之差餉及地租由甲方承擔。
3. 管理費由 {management_fee_payer} 承擔。

第七條 終止合約
1. 任何一方如需提前終止本合約，須提前兩個月以書面通知對方。
2. 若乙方違反本合約任何條款，甲方有權即時終止本合約。

第八條 其他
1. 本合約一式兩份，甲乙雙方各執一份，具同等法律效力。
2. 本合約未盡事宜，由甲乙雙方協商處理。

甲方簽署：_________________      日期：_________________

乙方簽署：_________________      日期：_________________
"""


class LeaseGenerator:
    """租约文件生成器"""

    def generate(
        self,
        unit_number: str,
        tenant_name: str,
        lease_start: date,
        lease_end: date,
        monthly_rent: Decimal,
    ) -> Tuple[BytesIO, str]:
        """
        生成租约文件

        返回：(file_bytes_io, filename)
        """
        # 计算押金（两个月租金）
        deposit = monthly_rent * 2
        # 计算租期
        duration_months = self._calc_duration(lease_start, lease_end)

        # 填充模板
        content = LEASE_TEMPLATE.format(
            lease_date=date.today().strftime("%Y年%m月%d日"),
            tenant_name=tenant_name,
            unit_number=unit_number,
            lease_start=lease_start.strftime("%Y年%m月%d日"),
            lease_end=lease_end.strftime("%Y年%m月%d日"),
            lease_duration=f"{duration_months}個月",
            monthly_rent=f"{monthly_rent:,.2f}",
            deposit=f"{deposit:,.2f}",
            management_fee_payer="乙方" if "車位" not in unit_number else "甲方",
        )

        # 生成 .docx 文件
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # 设置默认字体
            style = doc.styles["Normal"]
            font = style.font
            font.name = "PMingLiU"  # 新細明體
            font.size = Pt(11)

            # 标题
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run("培英中學 - 俊傑花園租賃合約")
            run.bold = True
            run.font.size = Pt(16)

            # 分割线
            doc.add_paragraph("─" * 40)

            # 将模板内容按段落分割
            paragraphs = content.split("\n\n")
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue

                lines = para_text.split("\n")
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    p = doc.add_paragraph()
                    if line.startswith("第") and "條" in line[:5]:
                        run = p.add_run(line)
                        run.bold = True
                        run.font.size = Pt(12)
                    elif "簽署" in line:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p.add_run(line)
                    else:
                        p.add_run(line)

                # 段落间距
                doc.add_paragraph("")

            # 保存到 BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            filename = f"租約_{unit_number}_{tenant_name}_{date.today().strftime('%Y%m%d')}.docx"
            return buffer, filename

        except ImportError:
            logger.warning("python-docx 未安裝，生成純文本租約")
            # 回退到纯文本
            buffer = BytesIO()
            buffer.write(content.encode("utf-8"))
            buffer.seek(0)
            filename = f"租約_{unit_number}_{tenant_name}_{date.today().strftime('%Y%m%d')}.txt"
            return buffer, filename

    def _calc_duration(self, start: date, end: date) -> int:
        """计算租期月数"""
        months = (end.year - start.year) * 12 + (end.month - start.month)
        if end.day > start.day:
            months += 1
        return max(months, 1)


lease_generator = LeaseGenerator()
